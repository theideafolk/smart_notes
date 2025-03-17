from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
import PyPDF2
from docx import Document
import io
import logging
from typing import Dict

app = FastAPI()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://localhost:5174"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract text from a PDF file content.
    """
    try:
        pdf_file = io.BytesIO(file_content)
        reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        raise ValueError("Failed to extract text from PDF")

def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text from a DOCX file content.
    """
    try:
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text.strip()
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        raise ValueError("Failed to extract text from DOCX")

@app.post("/extract-text")
async def extract_text_from_document(file: UploadFile = File(...)) -> Dict[str, str]:
    """
    Extract text from a document file (PDF or DOCX).
    """
    try:
        # Verify file type
        content_type = file.content_type
        if content_type not in [
            "application/pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword"
        ]:
            raise HTTPException(
                status_code=400,
                detail="Only PDF and DOC/DOCX files are supported"
            )

        # Read the file content
        content = await file.read()
        
        try:
            # Extract text based on file type
            if content_type == "application/pdf":
                text = extract_text_from_pdf(content)
            else:  # DOC/DOCX
                text = extract_text_from_docx(content)
            
            if not text:
                return {"text": f"No text could be extracted from the {file.filename}"}
            
            return {"text": text}
            
        except ValueError as ve:
            logger.error(f"Error extracting text: {str(ve)}")
            raise HTTPException(
                status_code=500,
                detail=str(ve)
            )
        
    except Exception as e:
        logger.error(f"Error processing document: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Error processing document: {str(e)}"
        )
    finally:
        await file.close()

@app.get("/health")
async def health_check():
    """
    Health check endpoint.
    """
    return {"status": "healthy"} 